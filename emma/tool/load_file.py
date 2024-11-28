from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any
import re


class LoadWordDoc:
    def load(self, filepath):
        # Load the Word document
        doc = Document(filepath)
        # Extract text from the document
        self.content = []
        for para in doc.paragraphs:
            self.content.append(para.text)
        # Join the paragraphs into a single string
        self.content = '\n'.join(self.content)
        return self
    
    def split_content(self):
        # Split the content into pairs using regex to match the pattern "number、"
        pairs = re.split(r'\d+、', self.content)[1:]  # Skip the first empty split
        # Extract query and answer from each pair
        extracted_pairs = []
        for pair in pairs:
            lines = pair.strip().split('\n')
            if len(lines) >= 2:
                query = lines[0].strip()
                ans = lines[1].strip()
                extracted_pairs.append({'query': query, 'ans': ans})
        return extracted_pairs
    
def extract_text(element):
    """Extract text from a docx element, removing styles."""
    text = ''
    for child in element.iter():
        if child.tag == qn('w:t'):
            text += child.text
    return text

def extract_tables(doc: Document) -> List[Dict[str, Any]]:
    """Extract tables from a docx document and convert them to markdown."""
    tables = []
    table_id = 1
    for table in doc.tables:
        md_table = '| ' + ' | '.join(['---'] * len(table.columns)) + ' |\n'
        for row in table.rows:
            md_table += '| ' + ' | '.join(cell.text.strip() for cell in row.cells) + ' |\n'
        tables.append({
            'table_id': f'table_{table_id}',
            'content': md_table
        })
        table_id += 1
    return tables

def extract_figures(doc: Document) -> List[Dict[str, Any]]:
    """Extract figures from a docx document."""
    figures = []
    figure_id = 1
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            figures.append({
                'figure_id': f'figure_{figure_id}',
                'url': rel.target_ref
            })
            figure_id += 1
    return figures

def get_page_number(paragraph, doc_text):
    """Get the page number of a paragraph using docx2txt."""
    para_text = paragraph.text.strip()
    if not para_text:
        return None
    pages = doc_text.split('\f')
    for page_number, page in enumerate(pages, start=1):
        if para_text in page:
            return page_number
    return None

def load_paragraphs(doc: Document, doc_name: str, doc_text: str) -> List[Dict[str, Any]]:
    """Load paragraphs from a docx document, replacing tables and figures with IDs."""
    paragraphs = []
    table_id = 1
    figure_id = 1
    for para in doc.paragraphs:
        text = extract_text(para)
        if not text.strip():
            continue
        page_number = get_page_number(para, doc_text)
        tables = []
        figures = []
        for table in doc.tables:
            if table in para._element:
                text = text.replace(table.text, f'<table_{table_id}>')
                tables.append(f'table_{table_id}')
                table_id += 1
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref and rel.target_ref in para._element.xml:
                text = text.replace(rel.target_ref, f'<figure_{figure_id}>')
                figures.append(f'figure_{figure_id}')
                figure_id += 1
        paragraphs.append({
            'document_name': doc_name,
            'content': text,
            'page_number': page_number,
            'tables': tables if tables else None,
            'figures': figures if figures else None
        })
    return paragraphs


# class LoadRegulationDoc:
#     def __init__(self) -> None:
#         # random tmp dir, timestamp with a random number
#         self.tmp_dir = f'/tmp/valacy/{int(time.time())}_{random.randint(1000, 9999)}'
#         os.makedirs(self.tmp_dir, exist_ok=True)
    
#     def convert_docx_to_text(self, input_dir):
#         output_dir = self.tmp_dir
#         if not os.path.exists(output_dir):
#             os.makedirs(output_dir)
#         for subdir, _, files in os.walk(input_dir):
#             for file in files:
#                 if file.endswith('.docx'):
#                     input_path = os.path.join(subdir, file)
#                     output_path = os.path.join(output_dir, file.replace('.docx', '.txt'))
#                     with open(input_path, 'rb') as docx_file:
#                         result = mammoth.extract_raw_text(docx_file)
#                         with open(output_path, 'w', encoding='utf-8') as text_file:
#                             text_file.write(result.value)
    
#     def load(self, filepath: str, filetype='txt') -> Dict[str, Any]:
#         if filetype == 'docx':
#             doc = Document(filepath)
#             doc_name = filepath.split('/')[-1].split('.')[0]
#             doc_text = docx2txt.process(filepath)
#             paragraphs = load_paragraphs(doc, doc_name, doc_text)
#             tables = extract_tables(doc)
#             figures = extract_figures(doc)
#             return {
#                 'paragraphs': paragraphs,
#                 'tables': tables,
#                 'figures': figures
#             }


if __name__ == '__main__':
    # Example usage
    file_path = '3国家电网有限公司采购业务实施细则.doc'
    # pairs = LoadWordDoc().load(file_path).split_content()
    # for pair in pairs:
    #     print(f"Query: {pair['query']}")
    #     print(f"Answer: {pair['ans']}")
    #     print()
    loader = LoadRegulationDoc()
    result = loader.load(file_path)
    for paragraph in result['paragraphs']:
        with open('output.txt', 'a') as f:
            f.write(f"{paragraph['content']}\t{paragraph['page_number']}\n\n\n")
    for table in result['tables']:
        print(table)
    for figure in result['figures']:
        print(figure)